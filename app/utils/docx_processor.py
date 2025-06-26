import logging
from io import BytesIO
from typing import Tuple

from docx import Document

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """Process DOCX files and extract text content."""

    def extract_text_from_docx(self, docx_bytes: bytes) -> Tuple[str, float]:
        """
        Extract text from DOCX file and assess quality.

        Args:
            docx_bytes: DOCX file as bytes

        Returns:
            Tuple of (extracted_text, quality_score)
        """
        try:
            # Load the document from bytes
            doc = Document(BytesIO(docx_bytes))

            # Extract text from all paragraphs
            text_parts = []
            total_chars = 0

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
                    total_chars += len(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
                            total_chars += len(cell.text.strip())

            # Combine all text
            full_text = "\n".join(text_parts)

            # Calculate quality score based on text content
            # DOCX files typically have high quality since they contain native text
            quality_score = min(
                100.0, max(80.0, total_chars / 10.0)
            )  # Base 80% for DOCX

            logger.info(
                f"Extracted {total_chars} characters from DOCX with quality score: {quality_score:.1f}"
            )

            return full_text.strip(), quality_score

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise

    def extract_text_from_file(self, file_path: str) -> Tuple[str, float]:
        """
        Extract text from DOCX file path.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (extracted_text, quality_score)
        """
        try:
            with open(file_path, "rb") as f:
                docx_bytes = f.read()

            return self.extract_text_from_docx(docx_bytes)

        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise
